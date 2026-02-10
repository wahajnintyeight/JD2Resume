import re
import io
from typing import Any, Dict, List, Tuple, Optional
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


DEFAULT_TEMPLATE_PATH = "templates/DummyCV.docx"


def _strip_html_tags(html: str) -> str:
    return re.sub(r'<[^>]+>', '', html)


def _add_tab_stops_to_paragraph(paragraph: Paragraph) -> None:
    """Add proper tab stops to a paragraph for left-center-right alignment."""
    # Only add tab stops if the paragraph contains tab characters
    if '\t' not in paragraph.text:
        return
    
    pPr = paragraph._element.get_or_add_pPr()
    
    # Remove existing tabs element if present
    tabs = pPr.find(qn('w:tabs'))
    if tabs is not None:
        pPr.remove(tabs)
    
    # Create new tabs element
    tabs = OxmlElement('w:tabs')
    
    # Center tab at 3.25 inches (for middle content like title/role)
    tab1 = OxmlElement('w:tab')
    tab1.set(qn('w:val'), 'center')
    tab1.set(qn('w:pos'), str(int(3.25 * 1440)))  # 1440 twips per inch
    tabs.append(tab1)
    
    # Right tab at 6.5 inches (for dates on the right)
    tab2 = OxmlElement('w:tab')
    tab2.set(qn('w:val'), 'right')
    tab2.set(qn('w:pos'), str(int(6.5 * 1440)))
    tabs.append(tab2)
    
    pPr.append(tabs)


def _get_nested_value(data: Dict[str, Any], path: List[str]) -> Any:
    current = data
    for key in path:
        if isinstance(current, dict):
            current = current.get(key)
            if current is None:
                return None
        else:
            return None
    return current


def _replace_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str]) -> None:
    """Replace placeholders in a paragraph while preserving formatting."""
    if not paragraph.runs:
        return
    
    paragraph_text = paragraph.text
    if not any(key in paragraph_text for key in replacements):
        return

    for run in paragraph.runs:
        for placeholder, replacement in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)


def _replace_across_runs(paragraph: Paragraph, replacements: Dict[str, str]) -> None:
    """Handle placeholders that might be split across multiple runs."""
    full_text = paragraph.text
    
    for placeholder, replacement in replacements.items():
        if placeholder not in full_text:
            continue
        
        runs_text = [(run, run.text) for run in paragraph.runs]
        combined = ""
        for run, text in runs_text:
            combined += text
        
        if placeholder in combined:
            start_idx = combined.find(placeholder)
            end_idx = start_idx + len(placeholder)
            
            char_count = 0
            start_run_idx = None
            end_run_idx = None
            start_char = None
            end_char = None
            
            for i, (run, text) in enumerate(runs_text):
                run_start = char_count
                run_end = char_count + len(text)
                
                if start_run_idx is None and run_start <= start_idx < run_end:
                    start_run_idx = i
                    start_char = start_idx - run_start
                
                if run_start < end_idx <= run_end:
                    end_run_idx = i
                    end_char = end_idx - run_start
                    break
                
                char_count = run_end
            
            if start_run_idx is not None and end_run_idx is not None:
                if start_run_idx == end_run_idx:
                    run = paragraph.runs[start_run_idx]
                    run.text = run.text[:start_char] + replacement + run.text[end_char:]
                else:
                    paragraph.runs[start_run_idx].text = (
                        paragraph.runs[start_run_idx].text[:start_char] + replacement
                    )
                    for i in range(start_run_idx + 1, end_run_idx):
                        paragraph.runs[i].text = ""
                    paragraph.runs[end_run_idx].text = (
                        paragraph.runs[end_run_idx].text[end_char:]
                    )


def _find_section_boundaries(
    doc: Document, 
    start_marker: str, 
    end_marker: str
) -> Tuple[Optional[int], Optional[int]]:
    """Find start and end paragraph indices for a section."""
    start_idx = None
    end_idx = None
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if start_marker in text:
            start_idx = idx
        elif end_marker in text and start_idx is not None:
            end_idx = idx
            break
    
    return start_idx, end_idx


def _clone_paragraph(template_para: Paragraph, insert_before_element) -> Paragraph:
    """Clone a paragraph and insert it before the given element."""
    new_elem = deepcopy(template_para._element)
    insert_before_element.addprevious(new_elem)
    return new_elem


def _process_repeating_section(
    doc: Document,
    start_marker: str,
    end_marker: str,
    items: List[Dict],
    placeholder_map: Dict[str, str],
    bullet_field: Optional[str] = None
) -> None:
    """Process a repeating section by cloning template paragraphs for each item."""
    start_idx, end_idx = _find_section_boundaries(doc, start_marker, end_marker)
    
    if start_idx is None or end_idx is None:
        return
    
    if not items:
        for idx in range(end_idx, start_idx - 1, -1):
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)
        return
    
    # Store references to the actual paragraph elements, not indices
    start_para_element = doc.paragraphs[start_idx]._element
    end_para_element = doc.paragraphs[end_idx]._element
    
    template_paras = []
    for idx in range(start_idx + 1, end_idx):
        para = doc.paragraphs[idx]
        if para.text.strip():
            template_paras.append(para)
    
    bullet_template_indices = []
    if bullet_field:
        bullet_placeholder = placeholder_map.get(bullet_field, "")
        for i, para in enumerate(template_paras):
            if bullet_placeholder in para.text:
                bullet_template_indices.append(i)
    
    # Insert position starts at the end marker
    insert_position = end_para_element
    
    for item in items:
        descriptions = []
        if bullet_field and bullet_field in placeholder_map:
            desc_key = placeholder_map[bullet_field].replace("{{", "").replace("}}", "").lower()
            for key, val in item.items():
                if key.lower() == "description":
                    descriptions = val if isinstance(val, list) else [val]
                    break
        
        # Process template paragraphs in order
        for i, template_para in enumerate(template_paras):
            if i in bullet_template_indices and descriptions:
                # Insert bullet points in order
                for desc in descriptions:
                    if desc:
                        new_elem = _clone_paragraph(template_para, insert_position)
                        for para in doc.paragraphs:
                            if para._element == new_elem:
                                clean_desc = _strip_html_tags(desc)
                                replacements = {placeholder_map[bullet_field]: clean_desc}
                                _replace_in_paragraph(para, replacements)
                                _replace_across_runs(para, replacements)
                                break
            elif i not in bullet_template_indices:
                new_elem = _clone_paragraph(template_para, insert_position)
                
                for para in doc.paragraphs:
                    if para._element == new_elem:
                        replacements = {}
                        for field, placeholder in placeholder_map.items():
                            if field == bullet_field:
                                continue
                            value = item.get(field, "")
                            if isinstance(value, list):
                                value = ", ".join(str(v) for v in value)
                            replacements[placeholder] = str(value) if value else ""
                        
                        _replace_in_paragraph(para, replacements)
                        _replace_across_runs(para, replacements)
                        
                        # Add tab stops for proper alignment if paragraph contains tabs
                        _add_tab_stops_to_paragraph(para)
                        break
    
    # Remove template paragraphs and markers using element references
    # Collect all elements to remove first
    elements_to_remove = [start_para_element, end_para_element]
    for template_para in template_paras:
        elements_to_remove.append(template_para._element)
    
    # Remove all collected elements
    for element in elements_to_remove:
        element.getparent().remove(element)


def _process_skills_section(doc: Document, additional: Dict) -> None:
    """Process skills section - replace skill placeholders with actual data."""
    start_idx, end_idx = _find_section_boundaries(
        doc, "{{SKILLS_START}}", "{{SKILLS_END}}"
    )
    
    if start_idx is None or end_idx is None:
        return
    
    tech_skills = additional.get("technicalSkills", [])
    
    backend_kw = ['laravel', 'nest', 'node', 'express', 'react', 'jquery', 
                  'bootstrap', 'django', 'flask', 'spring']
    db_kw = ['sql', 'postgres', 'mongo', 'redis', 'elastic', 'mysql']
    cloud_kw = ['aws', 'docker', 'gcp', 'azure', 'linux', 'git', 'rabbit', 
                'kubernetes', 'lambda', 's3']
    lang_kw = ['php', 'javascript', 'typescript', 'python', 'go', 'c#', 
               'java', 'ruby', 'html', 'css', 'sql']
    
    languages = []
    backend = []
    databases = []
    cloud = []
    
    for skill in tech_skills:
        s_lower = skill.lower()
        categorized = False
        
        for kw in lang_kw:
            if kw in s_lower:
                languages.append(skill)
                categorized = True
                break
        
        if categorized:
            continue
            
        for kw in db_kw:
            if kw in s_lower:
                databases.append(skill)
                categorized = True
                break
        
        if categorized:
            continue
            
        for kw in cloud_kw:
            if kw in s_lower:
                cloud.append(skill)
                categorized = True
                break
        
        if not categorized:
            backend.append(skill)
    
    replacements = {
        "PHP, JavaScript, Go, TypeScript, Python, C#, HTML5, CSS3, SQL": 
            ", ".join(languages) if languages else "N/A",
        "Laravel, NestJs, NodeJs, ExpressJS, ReactJS, jQuery, Bootstrap": 
            ", ".join(backend) if backend else "N/A",
        "PostgreSQL, Redis, Elasticsearch, MongoDB, MySQL": 
            ", ".join(databases) if databases else "N/A",
        "AWS (Lambda, S3), Docker, RabbitMQ, GCP, Azure, Linux, Git": 
            ", ".join(cloud) if cloud else "N/A",
    }
    
    for idx in range(start_idx + 1, end_idx):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            for old_text, new_text in replacements.items():
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
    
    for marker in ["{{SKILLS_START}}", "{{SKILLS_END}}"]:
        for para in list(doc.paragraphs):
            if marker in para.text:
                para._element.getparent().remove(para._element)
                break


def _process_summary_section(doc: Document, summary: str) -> None:
    """Process summary section."""
    start_idx, end_idx = _find_section_boundaries(
        doc, "{{SUMMARY_START}}", "{{SUMMARY_END}}"
    )
    
    if start_idx is None or end_idx is None:
        return
    
    if not summary or not summary.strip():
        for idx in range(end_idx, start_idx - 1, -1):
            doc.paragraphs[idx]._element.getparent().remove(
                doc.paragraphs[idx]._element
            )
        return
    
    for idx in range(start_idx + 1, end_idx):
        para = doc.paragraphs[idx]
        if "{{SUMMARY}}" in para.text:
            _replace_in_paragraph(para, {"{{SUMMARY}}": summary})
            _replace_across_runs(para, {"{{SUMMARY}}": summary})
    
    for marker in ["{{SUMMARY_START}}", "{{SUMMARY_END}}"]:
        for para in list(doc.paragraphs):
            if marker in para.text:
                para._element.getparent().remove(para._element)
                break


def _process_simple_placeholders(doc: Document, resume_data: Dict[str, Any]) -> None:
    """Replace simple placeholders."""
    personal_info = resume_data.get("personalInfo", {})
    
    replacements = {
        "{{NAME}}": personal_info.get("name", ""),
        "{{DESIGNATION_TITLE}}": personal_info.get("title", ""),
        "{{TITLE}}": personal_info.get("title", ""),
        "{{EMAIL}}": personal_info.get("email", ""),
        "{{PHONE}}": personal_info.get("phone", ""),
        "{{LOCATION}}": personal_info.get("location", ""),
        "{{LINKEDIN_URL}}": personal_info.get("linkedin", ""),
        "{{LINKEDIN}}": personal_info.get("linkedin", ""),
        "{{WEBSITE}}": personal_info.get("website", ""),
        "{{GITHUB}}": personal_info.get("github", ""),
    }
    
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)
        _replace_across_runs(para, replacements)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)
                    _replace_across_runs(para, replacements)


def generate_resume_from_template(
    resume_data: Dict[str, Any],
    template_path: str = DEFAULT_TEMPLATE_PATH
) -> bytes:
    """Generate a DOCX resume from a template file."""
    doc = Document(template_path)
    
    _process_simple_placeholders(doc, resume_data)
    
    _process_summary_section(doc, resume_data.get("summary", ""))
    
    work_exp = resume_data.get("workExperience", [])
    _process_repeating_section(
        doc,
        "{{EXPERIENCE_START}}",
        "{{EXPERIENCE_END}}",
        work_exp,
        {
            "company": "{{EXP_COMPANY}}",
            "title": "{{EXP_TITLE}}",
            "location": "{{EXP_LOCATION}}",
            "years": "{{EXP_DATES}}",
            "description": "{{EXP_BULLETS}}",
        },
        bullet_field="description"
    )
    
    _process_skills_section(doc, resume_data.get("additional", {}))
    
    projects = resume_data.get("personalProjects", [])
    proj_items = []
    for proj in projects:
        role_parts = proj.get("role", "").split(",")
        proj_items.append({
            "name": proj.get("name", ""),
            "tool1": role_parts[0].strip() if len(role_parts) > 0 else "",
            "tool2": role_parts[1].strip() if len(role_parts) > 1 else "",
            "toolx": ", ".join(p.strip() for p in role_parts[2:]) if len(role_parts) > 2 else "",
            "location": proj.get("location", ""),
            "years": proj.get("years", ""),
            "description": proj.get("description", []),
        })
    
    _process_repeating_section(
        doc,
        "{{PROJECTS_START}}",
        "{{PROJECTS_END}}",
        proj_items,
        {
            "name": "{{PROJECT_TITLE}}",
            "tool1": "{{PROJECT_TOOL1}}",
            "tool2": "{{PROJECT_TOOL2}}",
            "toolx": "{{PROJECT_TOOLX}}",
            "location": "{{PROJECT_LOCATION}}",
            "years": "{{PROJECT_DATES}}",
            "description": "{{PROJECT_BULLETS}}",
        },
        bullet_field="description"
    )
    
    education = resume_data.get("education", [])
    _process_repeating_section(
        doc,
        "{{EDUCATION_START}}",
        "{{EDUCATION_END}}",
        education,
        {
            "institution": "{{EDUCATION_INSTITUTE_NAME}}",
            "degree": "{{EDUCATION_DEGREE_NAME}}",
            "location": "{{EDUCATION_LOCATION}}",
            "years": "{{EDUCATION_DATES}}",
        }
    )
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()


def generate_resume_docx(resume_data: Dict[str, Any]) -> bytes:
    """Generate DOCX using default template."""
    return generate_resume_from_template(resume_data, DEFAULT_TEMPLATE_PATH)