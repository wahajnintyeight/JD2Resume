"""DOCX resume generation service."""

import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Any


def generate_resume_docx(resume_data: dict[str, Any]) -> bytes:
    """Generate a DOCX file from resume data.
    
    Args:
        resume_data: Dictionary containing resume sections
        
    Returns:
        Bytes of the generated DOCX file
    """
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Extract data
    personal_info = resume_data.get("personalInfo", {})
    summary = resume_data.get("summary", "")
    work_experience = resume_data.get("workExperience", [])
    education = resume_data.get("education", [])
    projects = resume_data.get("personalProjects", [])
    additional = resume_data.get("additional", {})
    
    # Header - Name
    if personal_info.get("name"):
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal_info["name"].upper())
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.name = "Times New Roman"
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    if personal_info.get("title"):
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(personal_info["title"])
        title_run.font.size = Pt(11)
        title_run.font.name = "Times New Roman"
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Line
    contact_parts = []
    if personal_info.get("location"):
        contact_parts.append(personal_info["location"])
    if personal_info.get("phone"):
        contact_parts.append(personal_info["phone"])
    if personal_info.get("email"):
        contact_parts.append(personal_info["email"])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.name = "Times New Roman"
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # LinkedIn
    if personal_info.get("linkedin"):
        linkedin_para = doc.add_paragraph()
        linkedin_label = linkedin_para.add_run("LinkedIn: ")
        linkedin_label.font.size = Pt(10)
        linkedin_label.font.name = "Times New Roman"
        
        linkedin_url = personal_info["linkedin"]
        if not linkedin_url.startswith("http"):
            linkedin_url = f"https://{linkedin_url}"
        linkedin_run = linkedin_para.add_run(linkedin_url)
        linkedin_run.font.size = Pt(10)
        linkedin_run.font.name = "Times New Roman"
        linkedin_run.underline = True
        linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Spacer
    doc.add_paragraph()
    
    # Summary
    if summary:
        _add_section_header(doc, "SUMMARY")
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.size = Pt(10.5)
        summary_run.font.name = "Times New Roman"
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Work Experience
    if work_experience:
        _add_section_header(doc, "WORK EXPERIENCE")
        for exp in work_experience:
            # Company and Location row
            exp_para = doc.add_paragraph()
            
            company_run = exp_para.add_run(exp.get("company", ""))
            company_run.bold = True
            company_run.font.size = Pt(11)
            company_run.font.name = "Times New Roman"
            
            if exp.get("location"):
                exp_para.add_run("\t")  # Tab
                location_run = exp_para.add_run(exp["location"])
                location_run.font.size = Pt(10)
                location_run.font.name = "Times New Roman"
                # Align right using tab stops would be better, but for simplicity:
                # We'll just let it flow
            
            # Title and Date row
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(exp.get("title", ""))
            title_run.italic = True
            title_run.font.size = Pt(10)
            title_run.font.name = "Times New Roman"
            
            if exp.get("years"):
                title_para.add_run("\t")
                years_run = title_para.add_run(exp["years"])
                years_run.font.size = Pt(10)
                years_run.font.name = "Times New Roman"
            
            # Description bullets
            for desc in exp.get("description", []):
                if desc:
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_run = bullet_para.add_run(_strip_html_tags(desc))
                    bullet_run.font.size = Pt(10.5)
                    bullet_run.font.name = "Times New Roman"
    
    # Skills
    technical_skills = additional.get("technicalSkills", [])
    languages = additional.get("languages", [])
    
    if technical_skills or languages:
        _add_section_header(doc, "SKILLS")
        
        if technical_skills:
            skills_para = doc.add_paragraph()
            label_run = skills_para.add_run("Languages: ")
            label_run.bold = True
            label_run.font.size = Pt(10.5)
            label_run.font.name = "Times New Roman"
            
            skills_run = skills_para.add_run(", ".join(technical_skills))
            skills_run.font.size = Pt(10.5)
            skills_run.font.name = "Times New Roman"
        
        if languages:
            lang_para = doc.add_paragraph()
            label_run = lang_para.add_run("Languages: ")
            label_run.bold = True
            label_run.font.size = Pt(10.5)
            label_run.font.name = "Times New Roman"
            
            lang_run = lang_para.add_run(", ".join(languages))
            lang_run.font.size = Pt(10.5)
            lang_run.font.name = "Times New Roman"
    
    # Projects
    if projects:
        _add_section_header(doc, "PROJECTS")
        for project in projects:
            # Project name and dates
            proj_para = doc.add_paragraph()
            name_run = proj_para.add_run(project.get("name", ""))
            name_run.bold = True
            name_run.font.size = Pt(11)
            name_run.font.name = "Times New Roman"
            
            if project.get("years"):
                proj_para.add_run("\t")
                years_run = proj_para.add_run(project["years"])
                years_run.font.size = Pt(10)
                years_run.font.name = "Times New Roman"
            
            # Tech stack
            if project.get("role"):
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run(project["role"])
                tech_run.italic = True
                tech_run.font.size = Pt(10)
                tech_run.font.name = "Times New Roman"
            
            # Description bullets
            for desc in project.get("description", []):
                if desc:
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_run = bullet_para.add_run(_strip_html_tags(desc))
                    bullet_run.font.size = Pt(10.5)
                    bullet_run.font.name = "Times New Roman"
    
    # Certifications
    certifications = additional.get("certificationsTraining", [])
    if certifications:
        _add_section_header(doc, "CERTIFICATIONS")
        for cert in certifications:
            cert_para = doc.add_paragraph(style='List Bullet')
            cert_run = cert_para.add_run(cert)
            cert_run.font.size = Pt(10.5)
            cert_run.font.name = "Times New Roman"
    
    # Awards
    awards = additional.get("awards", [])
    if awards:
        _add_section_header(doc, "AWARDS")
        for award in awards:
            award_para = doc.add_paragraph(style='List Bullet')
            award_run = award_para.add_run(award)
            award_run.font.size = Pt(10.5)
            award_run.font.name = "Times New Roman"
    
    # Education
    if education:
        _add_section_header(doc, "EDUCATION")
        for edu in education:
            # Institution and location
            edu_para = doc.add_paragraph()
            inst_run = edu_para.add_run(edu.get("institution", ""))
            inst_run.bold = True
            inst_run.font.size = Pt(11)
            inst_run.font.name = "Times New Roman"
            
            if edu.get("location"):
                edu_para.add_run("\t")
                loc_run = edu_para.add_run(edu["location"])
                loc_run.font.size = Pt(10)
                loc_run.font.name = "Times New Roman"
            
            # Degree and dates
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.get("degree", ""))
            degree_run.italic = True
            degree_run.font.size = Pt(10)
            degree_run.font.name = "Times New Roman"
            
            if edu.get("years"):
                degree_para.add_run("\t")
                years_run = degree_para.add_run(edu["years"])
                years_run.font.size = Pt(10)
                years_run.font.name = "Times New Roman"
            
            if edu.get("description"):
                desc_para = doc.add_paragraph()
                desc_run = desc_para.add_run(edu["description"])
                desc_run.font.size = Pt(10.5)
                desc_run.font.name = "Times New Roman"
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()


def _add_section_header(doc: Document, title: str) -> None:
    """Add a section header with underline."""
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(title)
    header_run.bold = True
    header_run.font.size = Pt(11)
    header_run.font.name = "Times New Roman"
    header_run.font.all_caps = True
    
    # Add underline using paragraph border (simplified - just text for now)
    # Word borders are complex, so we'll use a separator line
    doc.add_paragraph("_" * 60)


def _strip_html_tags(html: str) -> str:
    """Strip HTML tags from text."""
    import re
    clean = re.sub(r'<[^>]+>', '', html)
    return clean
